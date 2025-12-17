import streamlit as st
import json
import re
import time
import pdfplumber
from docx import Document
from io import BytesIO
from datetime import datetime
from groq import Groq

from pdf2image import convert_from_bytes

from PIL import Image, ImageOps
import pytesseract
import os

GROQ_API_KEY = os.environ["GROQ_API_KEY"]

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

st.set_page_config(
    page_title="GLR Pipeline - Insurance Template Filler",
    page_icon="🏠",
    layout="wide",
    initial_sidebar_state="collapsed"
)

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    
    .main { 
        font-family: 'Inter', sans-serif;
        background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%);
    }
    
    .stApp {
        background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%);
    }
    
    /* Header styling */
    .header-container {
        background: linear-gradient(135deg, #1e293b 0%, #334155 100%);
        padding: 3rem 2rem; 
        border-radius: 1.5rem; 
        margin-bottom: 2rem; 
        color: white; 
        text-align: center;
        box-shadow: 0 20px 40px rgba(30, 41, 59, 0.3);
    }
    
    .header-title {
        font-size: 3rem; 
        font-weight: 700; 
        margin-bottom: 1rem;
        background: linear-gradient(135deg, #60a5fa, #a78bfa);
        -webkit-background-clip: text; 
        -webkit-text-fill-color: transparent; 
        background-clip: text;
    }
    
    .header-subtitle { 
        font-size: 1.3rem; 
        opacity: 0.9; 
        font-weight: 400; 
    }
    
    /* Card styling */
    .content-card {
        background: white;
        padding: 2rem;
        border-radius: 1.5rem;
        box-shadow: 0 10px 30px rgba(0, 0, 0, 0.1);
        margin-bottom: 2rem;
        border: 1px solid #e2e8f0;
    }
    
    /* Upload zones */
    .upload-zone {
        background: linear-gradient(135deg, #f8fafc, #f1f5f9);
        border: 3px dashed #cbd5e1;
        border-radius: 1.5rem;
        padding: 2rem;
        text-align: center;
        margin: 1rem 0;
        transition: all 0.3s ease;
    }
    
    .upload-zone:hover {
        border-color: #3b82f6;
        background: linear-gradient(135deg, #eff6ff, #dbeafe);
    }
    
    .upload-zone.success {
        border-color: #10b981;
        background: linear-gradient(135deg, #ecfdf5, #d1fae5);
        border-style: solid;
    }
    
    /* Field display */
    .field-tag {
        background: linear-gradient(135deg, #1e293b, #334155);
        color: white;
        padding: 1rem;
        border-radius: 1rem;
        text-align: center;
        font-weight: 500;
        font-family: 'Monaco', 'Menlo', monospace;
        font-size: 0.9rem;
        margin: 0.5rem;
        box-shadow: 0 4px 15px rgba(30, 41, 59, 0.3);
    }
    
    /* Results styling */
    .results-container {
        background: white;
        border-radius: 1.5rem;
        box-shadow: 0 10px 30px rgba(0, 0, 0, 0.1);
        border: 1px solid #e2e8f0;
        margin: 1.5rem 0;
        overflow: hidden;
    }
    
    .results-header {
        background: linear-gradient(135deg, #f8fafc, #f1f5f9);
        padding: 1.5rem;
        border-bottom: 1px solid #e2e8f0;
        font-weight: 600;
        color: #1e293b;
        font-size: 1.1rem;
    }
    
    .results-row {
        padding: 1.5rem;
        border-bottom: 1px solid #f1f5f9;
        display: grid;
        grid-template-columns: 1fr 2fr;
        gap: 1rem;
    }
    
    .results-row:last-child {
        border-bottom: none;
    }
    
    .field-name {
        font-family: 'Monaco', 'Menlo', monospace;
        font-weight: 600;
        color: #475569;
        font-size: 0.9rem;
    }
    
    .field-value {
        color: #1e293b;
        font-weight: 500;
        font-size: 1rem;
    }
    
    /* Success styling */
    .success-container {
        background: linear-gradient(135deg, #ecfdf5, #d1fae5);
        border: 3px solid #10b981;
        border-radius: 1.5rem;
        padding: 3rem 2rem;
        text-align: center;
        margin: 2rem 0;
    }
    
    .success-title {
        font-size: 2.5rem;
        font-weight: 700;
        color: #059669;
        margin-bottom: 1rem;
    }
    
    .success-subtitle {
        color: #059669;
        font-size: 1.2rem;
        font-weight: 500;
    }
    
    /* Processing animation */
    .processing-container {
        text-align: center;
        padding: 3rem 2rem;
        background: white;
        border-radius: 1.5rem;
        box-shadow: 0 10px 30px rgba(0, 0, 0, 0.1);
        border: 1px solid #e2e8f0;
        margin: 2rem 0;
    }
    
    .processing-spinner {
        font-size: 4rem;
        margin-bottom: 2rem;
        animation: spin 2s linear infinite;
    }
    
    @keyframes spin {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }
    
    /* Button styling */
    .stButton > button {
        background: linear-gradient(135deg, #3b82f6, #2563eb) !important;
        color: white !important;
        border: none !important;
        border-radius: 1rem !important;
        padding: 0.75rem 2rem !important;
        font-weight: 600 !important;
        font-size: 1rem !important;
        transition: all 0.3s ease !important;
        box-shadow: 0 4px 15px rgba(59, 130, 246, 0.3) !important;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 8px 25px rgba(59, 130, 246, 0.4) !important;
    }
    
    .stDownloadButton > button {
        background: linear-gradient(135deg, #10b981, #059669) !important;
        color: white !important;
        border: none !important;
        border-radius: 1rem !important;
        padding: 0.75rem 1.5rem !important;
        font-weight: 600 !important;
        transition: all 0.3s ease !important;
        box-shadow: 0 4px 15px rgba(16, 185, 129, 0.3) !important;
    }
    
    .stDownloadButton > button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 8px 25px rgba(16, 185, 129, 0.4) !important;
    }
    
    /* Hide Streamlit elements */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    .stDeployButton {display: none;}
    
    /* Container styling */
    .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
        max-width: 1200px;
    }
</style>
""", unsafe_allow_html=True)

# Initialize Groq client
@st.cache_resource
def init_groq_client():
    api_key = GROQ_API_KEY
    if not api_key:
        st.error("Please set GROQ_API_KEY in Streamlit secrets")
        st.stop()
    return Groq(api_key=api_key)

# Session state initialization
if 'current_step' not in st.session_state:
    st.session_state.current_step = 1
if 'uploaded_template' not in st.session_state:
    st.session_state.uploaded_template = None
if 'uploaded_reports' not in st.session_state:
    st.session_state.uploaded_reports = []
if 'template_fields' not in st.session_state:
    st.session_state.template_fields = []
if 'extracted_data' not in st.session_state:
    st.session_state.extracted_data = {}
if 'processing_complete' not in st.session_state:
    st.session_state.processing_complete = False
if 'final_document' not in st.session_state:
    st.session_state.final_document = None

def extract_template_fields(docx_file):
    """Extract fields from DOCX template - fields are in [field] or {field} format"""
    try:
        doc = Document(docx_file)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
        
        fields = re.findall(r'(\[.*?\]|\{.*?\})', text)
        return list(set(fields))
    except Exception as e:
        st.error(f"Error extracting template fields: {str(e)}")
        return []

def extract_template_text(docx_file):
    """Extract text from DOCX template"""
    try:
        doc = Document(docx_file)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
        
        return text
    except Exception as e:
        st.error(f"Error extracting from template: {str(e)}")
        return []


def extract_pdf_text(pdf_file):
    """Extract text from PDF file page by page"""
    try:
        pdf_text = ""
        with pdfplumber.open(pdf_file) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    pdf_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
        return pdf_text
    except Exception as e:
        st.error(f"Error extracting PDF text: {str(e)}")
        return ""

def preprocess_image_for_ocr(img):
    # Convert to grayscale
    gray_image = ImageOps.grayscale(img)

    # Upscale by 2x
    scale_factor = 2
    resized_image = gray_image.resize(
        (gray_image.width * scale_factor, gray_image.height * scale_factor),
        resample=Image.LANCZOS
    )

    # Binarize (simple thresholding at 140)
    bw_image = resized_image.point(lambda x: 0 if x < 140 else 255, '1')

    return bw_image

def extract_pdf_text_with_ocr_fallback(pdf_file):
    try:
        # Try direct PDF text extraction first
        text = extract_pdf_text(pdf_file)
        pdf_file.seek(0)

        if len(text.strip()) > 250:  # Arbitrary threshold for "enough" text
            return text
        else:
            
            images = convert_from_bytes(pdf_file.read(), dpi=300)
            ocr_text = ""

            for i, img in enumerate(images):
                ocr_text += f"\n--- OCR Page {i + 1} ---\n"

                preprocessed_img = preprocess_image_for_ocr(img)

                # Tesseract config
                custom_config = r'--oem 3 --psm 6'

                page_text = pytesseract.image_to_string(preprocessed_img, config=custom_config)

                ocr_text += page_text

                # Optional: save preprocessed image for review
                # preprocessed_img.save(f"preprocessed_page_{i+1}.jpg")

            return ocr_text

    except Exception as e:
        st.error(f"OCR fallback failed: {str(e)}")
        return ""
    except Exception as e:
        st.error(f"OCR fallback failed: {str(e)}")
        return ""

def call_llm(template_text, template_fields, pdf_content):
    """Call Groq LLM to extract data"""
    try:
        client = init_groq_client()
        # solution = extract_template_text("Completed GLR Word Doc-ex2.docx")

        prompt = f"""
        You are an expert insurance document analyzer. Your task is to extract specific data from PDF reports to fill template fields. 

        DOCUMENT TYPE:
        (1) General Loss Reports Input Form — structured for submission to carriers, with sections like Claim Information, Insured Details, Property Address, Scope of Repair, and Damage Description.
        (2) Photo Documentation Reports — from inspection firms, focused on photo captions describing damage, location, cause, measurements, and observations. These include basic claim/policy/insured identifiers but lack the structured sections of General Loss Reports. 
        They provide supporting evidence for claim forms like the General Loss Report.
        All fields WILL be present in the PDF albeit with different names and variations. You will first understand the template thoroughly and then search the entire PDF content and extract the values.

        You MUST return ONLY a JSON object. The fields or placeholders in the template are always enclosed in **square [ ] or curly {{ }} brackets**.
        The keys in your output JSON **must exactly match the placeholders** as they appear in the template — including the brackets.
        ### TEMPLATE FIELDS TO EXTRACT
        {chr(10).join(template_fields)}

        ### FIELD MAPPING GUIDE
        - Date fields: Look for dates in MM/DD/YYYY, MM-DD-YYYY, or written formats
        - If DATE_LOSS is not explicitly provided, infer it as the earliest Date Taken from the Photo Sheets
        - Names: Look for contact names, estimator names, claim rep names
        - Numbers: Look for claim numbers, file numbers, policy numbers
        - Addresses: Look for street addresses, cities, states, zip codes
        - Phone/Email: Look for contact information
        - Coverage amounts: Look for dollar amounts, RCV values, ACV values
        - Company names: Look for insurance company names

        ### SEARCH STRATEGIES
        1. First pass: Look for exact matches
        2. Second pass: Look for partial matches and synonyms
        3. Third pass: Use context clues and proximity searches
        4. Consider that information might be in headers, footers, or photo captions
        5. Look for variations, abbreviations

        ### TEMPLATE CONTENT
        {template_text}

        ### PDF CONTENT TO SCAN
        {pdf_content}

        Begin analysis now:
        """
        chat_completion = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile", #"llama3-70b-8192",
            temperature=0.1
        )
        print("\nRESPONSE:", chat_completion)
        response_text = chat_completion.choices[0].message.content
        
        # Try to parse JSON from response
        try:
            json_start = response_text.find('{')
            json_end = response_text.rfind('}') + 1
            if json_start != -1 and json_end != 0:
                json_str = response_text[json_start:json_end]
                return json.loads(json_str)
            else:
                return json.loads(response_text)
        except json.JSONDecodeError:
            st.error("Error parsing LLM response as JSON")
            return {}
            
    except Exception as e:
        st.error(f"Error calling Groq API: {str(e)}")
        return {}

def populate_template(docx_file, field_data):
    """Replace template fields with extracted data"""
    try:
        doc = Document(docx_file)
        
        # Replace fields in paragraphs
        for paragraph in doc.paragraphs:
            for field, value in field_data.items():
                if value != "NOT_FOUND":
                    paragraph.text = paragraph.text.replace(field, str(value))
        
        # Replace fields in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for field, value in field_data.items():
                        if value != "NOT_FOUND":
                            cell.text = cell.text.replace(field, str(value))
        
        # Save to BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output
        
    except Exception as e:
        st.error(f"Error populating template: {str(e)}")
        return None

def render_header():
    st.markdown("""
    <div class="header-container">
        <div class="header-title">🏠 GLR Pipeline</div>
        <div class="header-subtitle">Professional Insurance Template Automation</div>
    </div>
    """, unsafe_allow_html=True)

def render_progress_bar():
    steps = ["📁 Upload", "🔍 Extract", "🤖 Process", "📝 Generate", "✅ Complete"]
    current = st.session_state.current_step
    
    progress_value = (current - 1) / (len(steps) - 1)
    st.progress(progress_value)
    
    cols = st.columns(5)
    for i, step in enumerate(steps, 1):
        with cols[i-1]:
            if i < current:
                st.markdown(f"<div style='text-align: center; color: #10b981; font-weight: 600;'>{step}</div>", unsafe_allow_html=True)
            elif i == current:
                st.markdown(f"<div style='text-align: center; color: #3b82f6; font-weight: 700; font-size: 1.1rem;'>{step}</div>", unsafe_allow_html=True)
            else:
                st.markdown(f"<div style='text-align: center; color: #94a3b8; font-weight: 500;'>{step}</div>", unsafe_allow_html=True)

def render_upload_step():
    st.markdown("## 📁 Upload Your Documents")
    st.markdown("Begin your automation journey by uploading your insurance template and inspection reports.")
    
    col1, col2 = st.columns(2, gap="large")
    
    with col1:
        st.markdown("### 📄 Insurance Template")
        
        if st.session_state.uploaded_template:
            st.markdown("""
            <div class="upload-zone success">
                <div style="font-size: 3rem; margin-bottom: 1rem;">✅</div>
                <div style="font-size: 1.2rem; font-weight: 600; color: #059669;">Template Uploaded Successfully!</div>
                <div style="color: #059669; margin-top: 0.5rem;">Ready for processing</div>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="upload-zone">
                <div style="font-size: 3rem; margin-bottom: 1rem;">📄</div>
                <div style="font-size: 1.2rem; font-weight: 600; color: #475569;">Upload DOCX Template</div>
                <div style="color: #64748b; margin-top: 0.5rem;">Drag and drop or click to browse</div>
            </div>
            """, unsafe_allow_html=True)

        template_file = st.file_uploader(
            "Choose DOCX Template",
            type=['docx'],
            key="template_uploader",
            accept_multiple_files=False,
            help="Upload a single DOCX template file with placeholder fields",
            label_visibility="collapsed"
        )
        
        if template_file:
            st.session_state.uploaded_template = template_file
    
    with col2:
        st.markdown("### 📁 Photo Reports")
        
        if st.session_state.uploaded_reports:
            st.markdown(f"""
            <div class="upload-zone success">
                <div style="font-size: 3rem; margin-bottom: 1rem;">✅</div>
                <div style="font-size: 1.2rem; font-weight: 600; color: #059669;">{len(st.session_state.uploaded_reports)} Reports Uploaded!</div>
                <div style="color: #059669; margin-top: 0.5rem;">Ready for analysis</div>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="upload-zone">
                <div style="font-size: 3rem; margin-bottom: 1rem;">📁</div>
                <div style="font-size: 1.2rem; font-weight: 600; color: #475569;">Upload PDF Reports</div>
                <div style="color: #64748b; margin-top: 0.5rem;">Multiple files supported</div>
            </div>
            """, unsafe_allow_html=True)
        
        report_files = st.file_uploader(
            "Choose PDF Reports",
            type=['pdf'],
            accept_multiple_files=True,
            key="reports_uploader",
            help="Upload multiple PDF inspection reports",
            label_visibility="collapsed"
        )
        
        if report_files:
            st.session_state.uploaded_reports = report_files
    
    if st.session_state.uploaded_template and st.session_state.uploaded_reports:
        st.markdown("---")
        col1, col2, col3 = st.columns([2, 1, 1])
        with col3:
            if st.button("Next: Extract Fields →", key="step1_next", use_container_width=True):
                st.session_state.template_text = extract_template_text(st.session_state.uploaded_template)
                st.session_state.template_fields = extract_template_fields(st.session_state.uploaded_template)
                st.session_state.current_step = 2
                st.rerun()

def render_extract_step():
    if not st.session_state.template_fields:
        st.error("❌ No template fields found. Please check your template format.")
        if st.button("← Back to Upload"):
            st.session_state.current_step = 1
            st.rerun()
        return
    
    st.markdown(f"## 🔍 Discovered {len(st.session_state.template_fields)} Template Fields")
    st.markdown("We've successfully identified the following placeholder fields in your template:")
    
    num_cols = 3
    cols = st.columns(num_cols)
    
    for i, field in enumerate(st.session_state.template_fields):
        with cols[i % num_cols]:
            st.markdown(f'<div class="field-tag">{field}</div>', unsafe_allow_html=True)
    
    st.markdown("---")
    col1, col2, col3 = st.columns([2, 1, 1])
    with col2:
        if st.button("← Back", use_container_width=True):
            st.session_state.current_step = 1
            st.rerun()
    with col3:
        if st.button("Next: AI Processing →", use_container_width=True):
            st.session_state.current_step = 3
            st.rerun()

def render_process_step():
    st.markdown("## 🤖 AI-Powered Data Extraction")
    st.markdown("Our advanced AI will analyze your documents and extract the relevant information.")
    
    if not st.session_state.processing_complete:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("🚀 Start AI Analysis", key="process_ai", use_container_width=True):
                # Show processing animation
                st.markdown("""
                <div class="processing-container">
                    <div class="processing-spinner">🤖</div>
                    <div style="font-size: 1.5rem; font-weight: 600; color: #1e293b; margin-bottom: 1rem;">AI Processing in Progress</div>
                    <div style="color: #64748b; font-size: 1rem;">Analyzing documents and extracting field data...</div>
                </div>
                """, unsafe_allow_html=True)
                
                with st.spinner("Processing documents..."):
                    # Extract text from all PDFs
                    all_pdf_content = ""
                    for pdf_file in st.session_state.uploaded_reports:
                        pdf_text = extract_pdf_text_with_ocr_fallback(pdf_file)
                        all_pdf_content += f"\n\n=== {pdf_file.name} ===\n{pdf_text}"

                    print("\nPDF ALL CONTENT: ", all_pdf_content)
                    print("\nTEMPLATE TEXT: ", st.session_state.template_text)
                    # Call LLM to extract data
                    extracted_data = call_llm(st.session_state.template_text, st.session_state.template_fields, all_pdf_content)
                    
                    if extracted_data:
                        st.session_state.extracted_data = extracted_data
                        st.session_state.processing_complete = True
                        st.success("✅ Data extraction completed successfully!")
                        time.sleep(1)
                        st.rerun()
                    else:
                        st.error("❌ Failed to extract data. Please try again.")
    else:
        # Show results
        st.markdown("### ✅ Extraction Complete!")
        
        st.markdown('<div class="results-container">', unsafe_allow_html=True)
        st.markdown('<div class="results-header">🎯 Extracted Field Data</div>', unsafe_allow_html=True)
        
        for field, value in st.session_state.extracted_data.items():
            color = "#059669" if value != "NOT_FOUND" else "#dc2626"
            st.markdown(f'''
            <div class="results-row">
                <div class="field-name">{field}</div>
                <div class="field-value" style="color: {color};">{value}</div>
            </div>
            ''', unsafe_allow_html=True)
        
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown("---")
        col1, col2, col3 = st.columns([2, 1, 1])
        with col2:
            if st.button("← Back", use_container_width=True):
                st.session_state.current_step = 2
                st.rerun()
        with col3:
            if st.button("Next: Generate Document →", use_container_width=True):
                st.session_state.current_step = 4
                st.rerun()

def render_generate_step():
    st.markdown("## 📝 Generate Final Document")
    st.markdown("Review the extraction results and generate your completed template.")
    
    found_fields = sum(1 for v in st.session_state.extracted_data.values() if v != "NOT_FOUND")
    total_fields = len(st.session_state.extracted_data)
    success_rate = int((found_fields / total_fields) * 100) if total_fields > 0 else 0
    
    # Show extraction summary with metrics using Streamlit's metric component
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Fields Extracted", f"{found_fields}/{total_fields}")
    with col2:
        st.metric("Success Rate", f"{success_rate}%")
    with col3:
        st.metric("Reports Processed", len(st.session_state.uploaded_reports))

    st.markdown("---")
    col1, col2, col3 = st.columns([2, 1, 1])
    with col2:
        if st.button("← Back", use_container_width=True):
            st.session_state.current_step = 3
            st.rerun()
    with col3:
        if st.button("🎯 Generate Document", use_container_width=True):    
            with st.spinner("Generating your completed document..."):
                final_doc = populate_template(st.session_state.uploaded_template, st.session_state.extracted_data)
                
                if final_doc:
                    st.session_state.final_document = final_doc
                    st.session_state.current_step = 5
                    st.rerun()
                else:
                    st.error("❌ Failed to generate document.")

def render_complete_step():
    st.markdown("""
    <div class="success-container">
        <div class="success-title">🎉 Document Generated Successfully!</div>
        <div class="success-subtitle">Your insurance template has been completed and is ready for download.</div>
    </div>
    """, unsafe_allow_html=True)
    
    # Show final metrics
    found_fields = sum(1 for v in st.session_state.extracted_data.values() if v != "NOT_FOUND")
    total_fields = len(st.session_state.extracted_data)
    success_rate = int((found_fields / total_fields) * 100) if total_fields > 0 else 0
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Fields Filled", f"{found_fields}/{total_fields}")
    with col2:
        st.metric("PDF Reports", len(st.session_state.uploaded_reports))
    with col3:
        st.metric("Success Rate", f"{success_rate}%")
    with col4:
        st.metric("Processing Time", "< 2min")
    
    st.markdown("---")
    st.markdown("## 📥 Download Your Results")
    
    col1, col2, col3 = st.columns(3, gap="large")
    
    with col1:
        if st.session_state.final_document:
            st.download_button(
                label="📄 Download Completed Template",
                data=st.session_state.final_document.getvalue(),
                file_name=f"completed_template_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    
    with col2:
        # Download extraction log
        log_data = {
            "timestamp": datetime.now().isoformat(),
            "template_fields": st.session_state.template_fields,
            "extracted_data": st.session_state.extracted_data,
            "success_rate": f"{success_rate}%",
            "reports_processed": len(st.session_state.uploaded_reports)
        }
        
        st.download_button(
            label="📊 Download Processing Log",
            data=json.dumps(log_data, indent=2),
            file_name=f"processing_log_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            mime="application/json",
            use_container_width=True
        )
    
    with col3:
        if st.button("🔄 Process Another Document", use_container_width=True):
            # Reset all session state
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()

def main():
    render_header()
    render_progress_bar()
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    if st.session_state.current_step == 1:
        render_upload_step()
    elif st.session_state.current_step == 2:
        render_extract_step()
    elif st.session_state.current_step == 3:
        render_process_step()
    elif st.session_state.current_step == 4:
        render_generate_step()
    elif st.session_state.current_step == 5:
        render_complete_step()

if __name__ == "__main__":
    main()
